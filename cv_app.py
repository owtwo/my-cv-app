from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('dp.jpg', width=Inches(2.0))

# name, phone number and email details
name = input('What is your name? ')
speak('Hello '+ name + ' how are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')
speak('What is your email?')
email = input('Enter your email: ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# About me
document.add_heading('About me')

speak('Tell me about your self?')
about_me = input('Tell me about your self? ')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company+' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ' '
)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experience = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company')
        from_date = input('From Date')
        to_date = input('To Date')

        p.add_run(company+' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company
        )
        p.add_run(experience_details)
    else:
        break

# Skills section
document.add_heading('Skills')
p = document.add_paragraph()

skill = input('Enter skill: ')
p.add_run(skill)
p.style = 'List Bullet'

# More skills
while True:
    has_more_skills = input('Do you have more skills? Yes or No')
    if has_more_skills.lower() == 'yes':
        p = document.add_paragraph()

        skill = input('Enter skill: ')
        p.add_run(skill)
        p.style = 'List Bullet'
    else:
        break


document.save('cv.docx')
